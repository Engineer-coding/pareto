"""DOCX reader using python-docx. Extracts paragraphs and detects headings."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document as DocxDocument

from pareto.ingestion.base import BaseReader
from pareto.ingestion.models import Document, DocumentFormat, StructuralHint


class DOCXReader(BaseReader):
    """
    Reads .docx files via python-docx.

    Heading detection: any paragraph whose style name starts with "Heading"
    (e.g. "Heading 1", "Heading 2") is recorded as a StructuralHint.
    """

    format = DocumentFormat.DOCX
    extensions = ("docx",)

    def read(self, path: str | Path) -> Document:
        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(f"DOCX file not found: {p}")

        doc = DocxDocument(str(p))

        # Walk paragraphs in order, building content + heading hints
        parts: list[str] = []
        hints: list[StructuralHint] = []
        cursor = 0

        for para in doc.paragraphs:
            text = para.text  # do NOT strip — paragraph order matters even if empty
            style_name = (para.style.name or "") if para.style else ""

            if text.strip() and style_name.startswith("Heading"):
                # Parse level from "Heading 2" etc.
                level = 1
                tail = style_name.replace("Heading", "").strip()
                if tail.isdigit():
                    level = int(tail)
                hints.append(
                    StructuralHint(
                        kind="heading",
                        level=level,
                        start_char=cursor,
                        end_char=cursor + len(text),
                        text=text.strip(),
                    )
                )

            parts.append(text)
            cursor += len(text) + 1  # +1 for the joining newline

        content = "\n".join(parts)

        # Try metadata title, else first H1 hint, else filename
        meta_title = None
        try:
            if doc.core_properties.title:
                meta_title = str(doc.core_properties.title).strip()
        except Exception:
            pass
        if not meta_title:
            for h in hints:
                if h.level == 1 and h.text:
                    meta_title = h.text
                    break
        title = meta_title or p.stem

        stat = p.stat()
        return Document(
            content=content,
            source=str(p.resolve()),
            format=self.format,
            title=title,
            structural_hints=hints,
            created_at=datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc),
            extra={
                "paragraph_count": len(doc.paragraphs),
                "heading_count": len(hints),
                "byte_size": stat.st_size,
            },
        )